#!/usr/bin/python
# -*- coding: utf-8 -*-


from WordStyleFunc import *
from EpibolyWorkTotalOnSystem import *
from docx.shared import Pt
from docx.shared import Cm  # 导入单位转换函数
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from WorkLogXls import *


def get_season_work_content(demand_no: str, logs: list):
    workload = 0
    tmp_content = []
    for log in logs:
        if demand_no == log['demand_no']:
            workload = workload + log['workload']
            if log['work_content'] not in tmp_content:
                tmp_content.append(log['work_content'])
    return tmp_content, workload


def get_season_log_name(person_name: str, year_and_season: str):
    return '附录5人月外包服务人员季度工作情况报告-成都思瑞奇信息产业有限公司-{}({}年{}季度).docx'.format(person_name, year_and_season[:4], year_and_season[4:])


def __set_cell_font(cell, **kwargs):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            font = run.font
            set_font_format(font, **kwargs)


def create_season_log(season_log: dict, epiboly_work: EpibolyWorkTotalOnSystem, person_cfg: json):
    year_and_season = season_log['year_and_season']
    person_name = season_log['person_name']
    company_name = person_cfg['person_info'].get(person_name).get('company')
    doc1 = init_season_doc(person_name, year_and_season[:4], year_and_season, company_name)

    table = doc1.add_table(rows=1, cols=4, style='Table Grid')
    table.style = 'Table Grid'
    table.width = Cm(15.77)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中

    table.rows[0].cells[0].text = '项目名称\r/涉及系统'
    table.rows[0].cells[1].text = '目前\r进度'
    table.rows[0].cells[2].text = '工作内容及工作时间\r（需列举出每项工作的人天天数）'
    table.rows[0].cells[3].text = '科技项目\r经理确认'

    table.rows[0].height = Cm(1.03)
    table.rows[0].cells[0].width = Cm(2.89)
    table.rows[0].cells[1].width = Cm(1.57)
    table.rows[0].cells[2].width = Cm(8.86)
    table.rows[0].cells[3].width = Cm(2.45)
    tot_workload = 0
    tmp_demand_no = ''
    demand_idx = 1
    for log in season_log['logs']:
        tot_workload = tot_workload + log['workload']
        if tmp_demand_no != log['demand_no']:
            tmp_demand_no = log['demand_no']
            table.add_row()
            table.rows[demand_idx].cells[0].text = log['demand_name'] + '/' + log['system_name']
            table.rows[demand_idx].cells[1].text = epiboly_work.demands[log['demand_no']]
            table.rows[demand_idx].cells[3].text = log['manager_name']
            log_contents, workload = get_season_work_content(log['demand_no'], season_log['logs'])
            for idx, log_content in enumerate(log_contents):
                p = table.rows[demand_idx].cells[2].add_paragraph(style='List Bullet')
                if idx == 0:
                    p = table.rows[demand_idx].cells[2].paragraphs[0]
                    p.style = 'List Bullet'

                run = p.add_run(log_content)
                f = run.font
                set_font_format(f, font_format={'size': 10.5, 'name': '仿宋'})
            if workload % 8 == 0:
                cast_tot = '以上工作累计花了{}个工作日。'.format(workload // 8)
            else:
                cast_tot = '以上工作累计花了{}个工作日。'.format(workload / 8)
            p = table.rows[demand_idx].cells[2].add_paragraph(style='List Continue')
            run = p.add_run(cast_tot)
            f = run.font
            set_font_format(f, font_format={'size': 10.5, 'name': '仿宋'})

            demand_idx = demand_idx + 1

    for row_idx in range(table.rows.__len__()):
        cells = table.row_cells(row_idx)
        for col_idx in range(cells.__len__()):
            cell = table.cell(row_idx, col_idx)
            # for cell in table.row_cells(row):
            if row_idx == 0:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                __set_cell_font(cell, font_format={'size': 14, 'name': '仿宋'})
            else:
                if col_idx == 0:
                    __set_cell_font(cell, font_format={'size': 10.5, 'name': '仿宋'})
                    table.rows[row_idx].cells[0].width = Cm(2.89)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                elif col_idx == 1:
                    __set_cell_font(cell, font_format={'size': 10.5, 'name': '仿宋'})
                    table.rows[row_idx].cells[1].width = Cm(1.57)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                elif col_idx == 2:
                    __set_cell_font(cell, font_format={'size': 10.5, 'name': '仿宋'})
                    table.rows[row_idx].cells[2].width = Cm(8.86)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif col_idx == 3:
                    __set_cell_font(cell, font_format={'size': 10.5, 'name': '仿宋'})
                    table.rows[row_idx].cells[3].width = Cm(2.45)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    end_text = '文件命名规则为“人月外包服务人员季度工作情况报告-所属公司-姓名(xx年xx季度)'
    p = doc1.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(end_text)
    f = run.font
    set_font_format(f, font_format={'size': 14, 'name': '仿宋'})
    pformat = p.paragraph_format
    paragraph_format2 = {'space_before': 0, 'space_after': 0, 'line_spacing': Pt(20)}
    set_paragraph_format(pformat, paragraph_format=paragraph_format2)

    log_out_dir = get_out_dir()
    doc1.save(os.path.join(log_out_dir, get_season_log_name(person_name, year_and_season)))
