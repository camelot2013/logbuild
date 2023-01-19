#!/usr/bin/python
# -*- coding: utf-8 -*-


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.font import Font
from docx.text.parfmt import ParagraphFormat
from docx.table import Table
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from typing import AnyStr


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_font_format(font: Font, **kwargs):
    """
    set font parameter
    Usage:
    set_font(font, font_param={'size': 22, 'name': '华文中宋'})
    设置字体: 华文中宋, 字体大小:22
    """
    if isinstance(font, Font):
        # list over all available tags
        for edge in ('font_format',):
            edge_data = kwargs.get(edge)
            if edge_data:
                # looks like order of attributes is important
                for key in ["size", "name", "bold"]:
                    if key in edge_data:
                        if key == 'size':
                            font.size = Pt(edge_data[key])
                        elif key == 'name':
                            font.name = edge_data[key]
                            font.element.rPr.rFonts.set(qn('w:eastAsia'), edge_data[key])  # 东亚地区的字符设置
                        elif key == 'bold':
                            font.bold = edge_data[key]


def set_paragraph_format(pformat: ParagraphFormat, **kwargs):
    if isinstance(pformat, ParagraphFormat):
        # list over all available tags
        for edge in ('paragraph_format',):
            edge_data = kwargs.get(edge)
            if edge_data:
                # looks like order of attributes is important
                for key in ["size", "name", "bold"]:
                    if key in edge_data:
                        if key == 'space_before':
                            pformat.space_before = Pt(edge_data[key])
                        elif key == 'space_after':
                            pformat.space_after = Pt(edge_data[key])
                        elif key == 'line_spacing_rule':
                            pformat.line_spacing_rule = edge_data[key]
                        elif key == 'line_spacing':
                            pformat.line_spacing = Pt(edge_data[key])


def merge(table: Table, row1, col1, row2, col2):
    if isinstance(table, Table):
        table.rows[row1].cells[col1].merge(table.rows[row2].cells[col2])


def __sort_key_demand(elem: dict):
    return elem['demand_no']


def init_season_doc(person_name: AnyStr, year_num: AnyStr, season_num: AnyStr, company_name: AnyStr) -> Document:
    # 新建空白文档
    doc1 = Document()

    # 设置中文内容的字体的方法
    head1_style = doc1.styles.add_style('Head1', WD_STYLE_TYPE.CHARACTER)
    head1_style.font.name = '黑体'
    head1_style.font.size = Pt(16)
    head1_style.font.color.rgb = RGBColor(0, 0, 0)
    head1_style.font.bold = False
    head1_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    head2_style = doc1.styles.add_style('Head2', WD_STYLE_TYPE.CHARACTER)
    head2_style.font.name = '华文中宋'
    head2_style.font.size = Pt(18)
    head2_style.font.color.rgb = RGBColor(0, 0, 0)
    head2_style.font.bold = False
    head2_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')

    head3_style = doc1.styles.add_style('Head3', WD_STYLE_TYPE.CHARACTER)
    head3_style.font.name = '仿宋'
    head3_style.font.size = Pt(14)
    head3_style.font.color.rgb = RGBColor(0, 0, 0)
    head3_style.font.bold = False
    head3_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    table_content_style = doc1.styles.add_style('table_content_style', WD_STYLE_TYPE.CHARACTER)
    table_content_style.font.name = '仿宋'
    table_content_style.font.size = Pt(14)
    table_content_style.font.color.rgb = RGBColor(0, 0, 0)
    table_content_style.font.bold = False
    table_content_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    p = doc1.add_heading(level=2)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run('附录5', style='Head1')
    f = run.font
    f.bold = False

    p = doc1.add_heading(level=2)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('人月外包服务人员季度工作情况报告', style='Head2')
    f = run.font
    f.bold = False

    p = doc1.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text1 = "({}年{}季度)".format(year_num, season_num)
    p.add_run(text1, style='Head3')

    p = doc1.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    text1 = "姓名：{}        公司名称：{}".format(person_name, company_name)
    p.add_run(text1, style='Head3')
    pformat = p.paragraph_format
    paragraph_format = {'space_before': 0, 'space_after': 0, 'line_spacing_rule': WD_LINE_SPACING.SINGLE}
    set_paragraph_format(pformat, paragraph_format=paragraph_format)

    return doc1