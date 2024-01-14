#!/usr/bin/python
# -*- coding: utf-8 -*-
import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from WordStyleFunc import set_paragraph_format, set_cell_font, merge
from WorkLogXls import read_cfg, get_out_dir


class CorpSeasonLog(object):
    @property
    def cfg(self):
        return self.__cfg

    @property
    def document(self):
        return self.__document

    def __init__(self, __year_num: str, __season_num: str, __corp_name: str):
        self.__document = Document()
        self.__cfg = read_cfg()
        self.__year_num = __year_num
        self.__season_num = __season_num
        self.__corp_name = __corp_name

    def init_document_head(self):
        # 设置中文内容的字体的方法
        head1_style = self.__document.styles.add_style('Head1', WD_STYLE_TYPE.CHARACTER)
        head1_style.font.name = '黑体'
        head1_style.font.size = Pt(18)
        head1_style.font.color.rgb = RGBColor(0, 0, 0)
        head1_style.font.bold = False
        head1_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        text_style = self.__document.styles.add_style('text', WD_STYLE_TYPE.CHARACTER)
        text_style.font.name = '仿宋'
        text_style.font.size = Pt(12)
        text_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        def __add_run(__text: str, __is_head: bool, __head_level, __style, __alignment, __paragraph_format):
            if __is_head:
                p = self.document.add_heading(level=__head_level)
            else:
                p = self.document.add_paragraph()
            p.add_run(__text, style=__style)
            p.alignment = __alignment
            pformat = p.paragraph_format
            set_paragraph_format(pformat, paragraph_format=__paragraph_format)

        text1 = '外包技术服务季度工作情况报告'
        paragraph_format = {'space_before': 0, 'space_after': 0, 'line_spacing_rule': WD_LINE_SPACING.SINGLE}
        __add_run(text1, True, 1, 'Head1', WD_PARAGRAPH_ALIGNMENT.CENTER, paragraph_format)

        text1 = f'（{self.__year_num}年{self.__season_num}季度）'
        __add_run(text1, False, 0, 'Head1', WD_PARAGRAPH_ALIGNMENT.CENTER, paragraph_format)

        text1 = f'公司名称：{self.__corp_name}           时间：{self.__year_num}年 {self.__season_num}季度'
        __add_run(text1, False, 0, 'text', WD_PARAGRAPH_ALIGNMENT.LEFT, paragraph_format)

    @staticmethod
    def __set_cell_content(__cell, __content: str, __alignment):
        __cell.text = __content
        __cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        __cell.paragraphs[0].alignment = __alignment
        set_cell_font(__cell, font_format={'size': 12, 'name': '仿宋'})

    def __fill_table_body(self, __table, corp_season_work: [dict]):
        def __get_current_index(__table):
            return __table.rows.__len__() - 1

        def int_or_float(v1: float):
            return int(v1) if str(v1).endswith(".0") else v1

        for person_season in corp_season_work:
            person_name = person_season.get('person_name')
            person_level = person_season.get('person_level')
            tot_workload = 0
            person_began_index = __get_current_index(__table) + 1
            for system_works in person_season.get('system'):
                system_name = system_works.get('system_name')
                system_began_index = __get_current_index(__table) + 1
                system_total_work = 0
                for __work in system_works.get('works'):
                    system_total_work = system_total_work + __work.get('workload')
                    tot_workload = tot_workload + __work.get('workload')
                    __table.add_row()
                    current_index = __get_current_index(__table)
                    __cell = __table.rows[current_index].cells[2]
                    self.__set_cell_content(__table.rows[current_index].cells[2], f'[{__work["month_request_no"]}]{__work["demand_name"]}', WD_PARAGRAPH_ALIGNMENT.LEFT)
                    workload = __work.get("workload")
                    self.__set_cell_content(__table.rows[current_index].cells[3], f'{int_or_float(workload)}人天;', WD_PARAGRAPH_ALIGNMENT.CENTER)
                    self.__set_cell_content(__table.rows[person_began_index].cells[0], f'{person_name}:{person_level}（成都）', WD_PARAGRAPH_ALIGNMENT.CENTER)

                    self.__set_cell_content(__table.rows[system_began_index].cells[1], f'{system_name}（{int_or_float(system_total_work)}人天）', WD_PARAGRAPH_ALIGNMENT.CENTER)
                current_index = __get_current_index(__table)
                merge(__table, system_began_index, 1, current_index, 1)

            __table.add_row()
            current_index = __get_current_index(__table)
            self.__set_cell_content(__table.rows[current_index].cells[1], '合计', WD_PARAGRAPH_ALIGNMENT.CENTER)

            self.__set_cell_content(__table.rows[current_index].cells[3], f'{int_or_float(tot_workload)}人天；', WD_PARAGRAPH_ALIGNMENT.CENTER)
            merge(__table, current_index, 1, current_index, 2)
            merge(__table, person_began_index, 0, current_index, 0)

    def create_corp_season_log(self, corp_season_work: [dict]):
        self.init_document_head()
        table = self.document.add_table(rows=1, cols=5, style='Table Grid')
        table.style = 'Table Grid'
        table.width = Cm(16.03)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中

        self.__set_cell_content(table.rows[0].cells[0], '服务人员\r岗位级别\r（工作地）', WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.__set_cell_content(table.rows[0].cells[1], '项目名称/涉及系统', WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.__set_cell_content(table.rows[0].cells[2], '需求名称', WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.__set_cell_content(table.rows[0].cells[3], '工作时间\r（需列举出每项工作的人天天数）', WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.__set_cell_content(table.rows[0].cells[4], '科技项目经理确认', WD_PARAGRAPH_ALIGNMENT.CENTER)

        table.rows[0].height = Cm(1.96)
        table.rows[0].cells[0].width = Cm(2.19)
        table.rows[0].cells[1].width = Cm(2.84)
        table.rows[0].cells[2].width = Cm(5.66)
        table.rows[0].cells[3].width = Cm(3.75)
        table.rows[0].cells[4].width = Cm(1.59)

        self.__fill_table_body(table, corp_season_work)

        log_out_dir = get_out_dir()
        self.document.save(os.path.join(log_out_dir, f'外包技术服务季度工作情况报告-{self.__corp_name}（{self.__year_num}年{self.__season_num}季度）.docx'))
