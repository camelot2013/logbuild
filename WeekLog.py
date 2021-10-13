#!/usr/bin/python
# -*- coding: utf-8 -*-


from WorkLogXls import *
from WordStyleFunc import set_font_format
from WordStyleFunc import set_paragraph_format
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm  # 导入单位转换函数
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def merge(table, row1, col1, row2, col2):
    table.rows[row1].cells[col1].merge(table.rows[row2].cells[col2])


def get_weeklog_filename(begin_date: str, end_date: str, person_name: str):
    return '人月外包服务人员个人周报-思瑞奇-{}({}-{}).docx'.format(person_name, begin_date, end_date)


# def walk_weeks(weeks: list):
#     for week in weeks:
#         create_weeklog(week)


# 检查工作日志内容是否包含有效内容
def check_work_log(work_log: dict):
    if 'work_date' not in work_log:
        print('日志内容不包含工作日期')
        return False
    if 'person_name' not in work_log:
        print('日志内容不包含人员名称')
        return False
    if 'demand_name' not in work_log:
        print('日志内容不包含需求项目名称')
        return False
    if 'manager_name' not in work_log:
        print('日志内容不包含项目经理')
        return False
    if 'month_request_no' not in work_log:
        print('日志内容不包人月申请编号')
        return False
    if 'workload' not in work_log:
        print('日志内容不包工作量')
        return False
    if 'work_content' not in work_log:
        print('日志内容不包工作内容')
        return False
    return True


def tabBgColor(table, row_idx, cols, colorStr):
    shading_list = locals()
    for i in range(cols):
        shading_list['shading_elm_' + str(i)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table.rows[row_idx].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])


def create_weeklog(work_log_week: list):
    # 传入参数合法性检查
    if work_log_week.__len__() > 7:
        print(work_log_week)
        print('一周的日志内容最多包含7天')
        return
    for work_log in work_log_week:
        if not check_work_log(work_log):
            return
    # 新建空白文档
    doc1 = Document()

    # 设置中文内容的字体的方法
    head1_style = doc1.styles.add_style('Head1', WD_STYLE_TYPE.CHARACTER)
    head1_style.font.name = '华文中宋'
    head1_style.font.size = Pt(22)
    head1_style.font.color.rgb = RGBColor(0, 0, 0)
    head1_style.font.bold = False

    head1_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
    head2_style = doc1.styles.add_style('Head2', WD_STYLE_TYPE.CHARACTER)
    head2_style.font.name = '楷体'
    head2_style.font.size = Pt(16)
    head2_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    text_style = doc1.styles.add_style('text', WD_STYLE_TYPE.CHARACTER)
    text_style.font.name = '仿宋'
    text_style.font.size = Pt(12)
    text_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    # 设置中文字体示例结束
    p = doc1.add_heading(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('人月外包服务人员工作周报', style='Head1')
    f = run.font
    f.bold = False

    p = doc1.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    begin_date = work_log_week[0]['work_date']
    end_date = work_log_week[-1]['work_date']
    person_name = work_log_week[0]['person_name']
    text1 = "( {} 年 {} 月 {} 日 ---- {} 年 {} 月 {} 日 )".format(begin_date[:4],
                                                            begin_date[4:6],
                                                            begin_date[6:],
                                                            end_date[:4],
                                                            end_date[4:6],
                                                            end_date[6:])
    run = p.add_run(text1, style='Head2')
    pformat = p.paragraph_format
    paragraph_format = {'space_before': 0, 'space_after': 0, 'line_spacing_rule': WD_LINE_SPACING.SINGLE}
    set_paragraph_format(pformat, paragraph_format=paragraph_format)

    p = doc1.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    text2 = '姓名：' + person_name + '                           所属公司：成都思瑞奇信息有限公司'
    run = p.add_run(text2, style='text')
    pformat = p.paragraph_format
    set_paragraph_format(pformat, paragraph_format=paragraph_format)

    table = doc1.add_table(rows=11, cols=6, style='Table Grid')
    table.style = 'Table Grid'
    table.width = Cm(15.18)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中

    table.rows[1].cells[1].text = '具体任务描述'
    table.rows[1].cells[2].text = '具体工作描述'
    table.rows[1].cells[3].text = '完成/正常/提前/延期'
    table.rows[1].cells[4].text = '工作进展程度及其他需要说明的内容'
    table.rows[2].cells[0].text = '星期一'
    table.rows[3].cells[0].text = '星期二'
    table.rows[4].cells[0].text = '星期三'
    table.rows[5].cells[0].text = '星期四'
    table.rows[6].cells[0].text = '星期五'
    table.rows[7].cells[0].text = '星期六'
    table.rows[8].cells[0].text = '星期日'
    table.rows[9].cells[0].text = '下周主要工作任务'
    # 背景色，根据需要调整，可参考站长之家选色 http://tool.chinaz.com/Tools/PageColor.aspx
    colorStr = 'D3D3D3'

    tabBgColor(table, 7, 6, colorStr)
    tabBgColor(table, 8, 6, colorStr)

    merge(table, 0, 0, 1, 0)
    merge(table, 0, 5, 1, 5)
    merge(table, 9, 0, 9, 5)
    merge(table, 10, 0, 10, 5)

    table.rows[0].height = Cm(0.5)
    table.rows[1].height = Cm(1.1)
    table.rows[2].height = Cm(2.5)
    table.rows[3].height = Cm(2.5)
    table.rows[4].height = Cm(2.5)
    table.rows[5].height = Cm(2.5)
    table.rows[6].height = Cm(2.5)
    table.rows[7].height = Cm(1.3)
    table.rows[8].height = Cm(1.3)
    table.rows[9].height = Cm(1.0)
    table.rows[10].height = Cm(2.5)
    table.rows[0].cells[0].text = '工作日'
    table.rows[0].cells[1].text = '项目名称'
    table.rows[0].cells[2].text = '工作描述'
    table.rows[0].cells[3].text = '状态'
    table.rows[0].cells[4].text = '说明'
    table.rows[0].cells[5].text = '项目负责人确认'
    for work_log in work_log_week:
        week_idx = work_log['week']
        table.rows[week_idx + 2].cells[1].text = work_log['demand_name']
        table.rows[week_idx + 2].cells[2].text = work_log['work_content']
        table.rows[week_idx + 2].cells[3].text = '正常'
        table.rows[week_idx + 2].cells[4].text = '完成'
        table.rows[week_idx + 2].cells[5].text = work_log['manager_name']

    table.rows[0].cells[0].width = Cm(1.57)
    table.rows[0].cells[1].width = Cm(2.64)
    table.rows[0].cells[2].width = Cm(4.07)
    table.rows[0].cells[3].width = Cm(1.98)
    table.rows[0].cells[4].width = Cm(2.91)
    table.rows[0].cells[5].width = Cm(2.06)

    for row_idx in range(table.rows.__len__()):
        cells = table.row_cells(row_idx)
        for col_idx in range(cells.__len__()):
            cell = table.cell(row_idx, col_idx)
            # for cell in table.row_cells(row):
            if row_idx in [2, 3, 4, 5, 6, 7, 8] and col_idx == 1:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            else:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraphs = cell.paragraphs
            # print('row_idx', row_idx, 'col_idx', col_idx)
            for paragraph in paragraphs:
                pformat = paragraph.paragraph_format
                set_paragraph_format(pformat, paragraph_format=paragraph_format)
                # print('set_paragraph_format')
                if row_idx in [2, 3, 4, 5, 6, 7, 8]:
                    if col_idx == 1:
                        pformat.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    elif col_idx in [2, 5]:
                        pformat.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    pformat.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for run in paragraph.runs:
                    font = run.font
                    # print('set_font')
                    set_font_format(font, font_format={'size': 12, 'name': '仿宋'})

    p = doc1.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pformat = p.paragraph_format
    set_paragraph_format(pformat, paragraph_format=paragraph_format)
    text2 = '注：文件命名规则为“人月外包服务人员个人周报-所属公司-姓名(YYYYMMDD-YYYYMMDD)”'
    p.add_run(text2, style='text')
    log_out_dir = get_out_dir()
    doc1.save(os.path.join(log_out_dir, get_weeklog_filename(begin_date, end_date, person_name)))
